"""生成项目技术总结 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_font(run, size=11, bold=False, color=None):
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(31, 78, 121))
    # 下划线
    run.font.underline = True
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(0, 70, 127))
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(68, 114, 196))
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent * 0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def bullet(text, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent * 0.3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def note(text):
    """灰色说明文字"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, color=(100, 100, 100))
    return p

# ── 标题 ──────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
r = title.add_run('Chef Vision 项目技术总结报告')
set_font(r, size=18, bold=True, color=(31, 56, 100))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(16)
r2 = sub.add_run('炒菜机器人食材温度实时监测系统')
set_font(r2, size=12, color=(80, 80, 80))

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
h1('一、食材区域追踪')

h2('1.1 采取的策略')
body('使用 SAM2 视频分割模型，以 100 帧（4 秒）为一批对食材区域进行自动追踪。系统初始化需要两步人工操作（仅需一次）：')
bullet('RGB 首帧标注：在第一帧手动标注食材前景点和背景点（LabelFirstFrame.py），生成 food_labels.json')
bullet('IR 锅区域圈选：在 IR 图像上手动标注锅的椭圆区域（cx/cy/rx/ry）和旋转轴坐标（wok_region.json）')
body('后续完全自动运行：每批用上一批末帧 mask 传递边界（carry_mask），IR 只做"守门"检查，不干预 SAM2 的视觉追踪过程。')

h2('1.2 存在的问题')
h3('SAM2 追踪精度不稳定')
body('改为全自动模式后，SAM2 靠视觉语义自由追踪，在翻炒、白烟、手入镜等场景下容易跑偏，批间误差可能逐步积累。')
body('根本原因：SAM2 是通用分割模型，非炒菜专用；手持拍摄带来的视角抖动、光线变化、白烟遮挡均会干扰语义判断。若设备固定在炒菜机器人上，摄像头与锅的相对位置固定、视角稳定、无手持抖动，上述问题大部分会自动消除。')

h2('1.3 针对性应对策略')
bullet('IR-IoU 门控：每批将 SAM2 mask 与 IR K-means 低温区计算 IoU，低于 15% 则强制重置，提前发现语义反转')
bullet('逐帧 IR-fix：当 mask 超过锅区域 50% 时，当帧立即用 IR K-means 重新生成 mask，不等批次结束')
bullet('面积三级检测：超过 wok 35%、骤降超 70%、绝对值低于 2% 三种情况均触发重置')
bullet('wok 中心动态追踪：每批用热环像素 Kasa 圆拟合法估算锅的几何圆心，修正手持抖动引起的位置漂移')
bullet('白烟/空锅冻结：检测到锅内 RGB 均匀高亮或极暗时冻结 mask，跳过本批补强')

# ══════════════════════════════════════════════════════════════════════════════
h1('二、三种测温方案')

h2('2.1 方案一：SAM2 Mask 追踪温度')
body('原理：将 SAM2 食材 mask 通过单应矩阵投影到 IR 坐标系，统计投影区域内的温度均值/最大值/最小值。')
h3('问题')
body('追踪精度直接影响测温精度；单应矩阵投影误差约 ±15px（手持拍摄时更大）。')
h3('应对')
body('IR-fix 和重置机制保证 mask 不会严重跑偏；设备固定后投影误差恒定可重新精确标定。')

h2('2.2 方案二：ROI 固定圆圈温度')
body('原理：在 RGB 上手动设定一个固定圆圈（圆心+半径），投影到 IR 坐标系后统计区域温度均值。')
h3('问题')
body('食材翻炒时可能偏移出圆圈，导致测温区域与实际食材位置不重合，读数偏差大。')
h3('应对')
body('适合食材始终在锅底中心区域的场景，与方案一配合使用互为验证；炒菜机器人固定设备后食材运动轨迹更可预测，ROI 有效性提升。')

h2('2.3 方案三：IR 自动分割（K-means）')
body('原理：直接在 IR 锅区域内做双峰 K-means 分类，低温类为食材、高温类为锅壁，取低温类均值作为食材温度。')
h3('问题')
body('IR 分辨率仅 192×256，空间定位粗；锅倾斜翻炒时锅内温度分布均匀（两峰差 < 30°C），无法区分食材和锅壁。')
h3('应对')
body('倾斜期间返回 NaN 跳过该帧，不污染温度曲线；与方案一形成高精度/低精度互补，IR 方案可作为 SAM2 方案的独立验证。')

h2('2.4 测温精度说明')
body('与实际测温针对比结果：')
bullet('最高温度：设备检测峰值与测温针误差在 5°C 以内，峰值读数可信')
bullet('平均温度：设备统计均值比测温针偏低约 10°C，原因是 mask 覆盖了部分温度较低的锅边食材和未完全受热区域，拉低了整体均值')

# ══════════════════════════════════════════════════════════════════════════════
h1('三、实时监测目标的瓶颈')

h2('3.1 问题：推理延迟，暂不满足硬实时')
body('当前每批（100 帧 = 4 秒视频）处理时间约 18 秒（帧抽取 ~5s + SAM2 推理 ~13s），处理速度慢于实时录制速度。')

h2('3.2 针对性应对策略')
bullet('并行流水线：采集线程与推理线程并行运行，采集第 N+1 批时同步推理第 N 批，实现近实时输出（滞后约 1 个批次）')
bullet('降低批次帧数：从 100 帧缩至 50 帧，单批延迟减半，代价是批间 carry_mask 传递频率加倍，轻微增加重置风险')
bullet('轻量化模型：用 SAM2-tiny 替代 SAM2-large，推理时间压缩约 50%，精度有所损失但速度大幅提升')
note('注：若设备固定后场景稳定，SAM2 追踪失控频率降低，可适当减少 IR-fix 等校验开销，进一步压缩延迟。')

# ══════════════════════════════════════════════════════════════════════════════
h1('四、生/熟/焦状态识别')

h2('4.1 问题：网络数据训练效果差，难以落地')

h3('原因一：数据质量差')
body('网络爬取的炒菜视频包含大量非纯炒菜内容（切菜、讲解、摆盘等），有效烹饪片段占比低，数据清洗成本高，标注工作量大。')

h3('原因二：域迁移问题')
body('网络视频经过后期调色处理，光线、色调、拍摄角度与炒菜机器人实际摄像头画面差距大，在网络数据上训练的 RGB 分类模型在实机上泛化能力差。')

h3('原因三：IR 数据缺乏')
body('红外烹饪视频在公开渠道几乎不存在，无法靠公开数据训练 IR 维度的分类模型，而 IR 信息对于焦糊检测尤为关键。')

h2('4.2 针对性应对策略')
body('最可行的路径：用炒菜机器人实际采集 RGB+IR 配对视频，人工标注关键状态时刻（食材入锅、变色成熟、出现焦糊），建立专用训练集。实机数据天然解决了光线、视角、色调差异，且 RGB 与 IR 严格同步，分类特征更可信，训练出的模型对本机场景精度最高。')
bullet('短期方案：用 IR 温度阈值规则做粗判断（如食材平均温度 > 180°C 持续 30s 判断为接近焦糊），无需训练数据')
bullet('中期方案：采集 50~100 组实机炒菜数据，人工标注，训练轻量分类网络（MobileNet/EfficientNet-lite）')
bullet('长期方案：结合 RGB 视觉特征（颜色变化）+ IR 温度特征（双峰分布变化）做多模态融合分类，提升各类菜品的泛化能力')

doc.add_paragraph()
note('本报告生成时间：2026年6月22日')

# ── 保存 ──────────────────────────────────────────────────────────────────────
out_path = r'D:\Chef_Vision\Chef_Vision_技术总结报告.docx'
doc.save(out_path)
print(f'已保存: {out_path}')
